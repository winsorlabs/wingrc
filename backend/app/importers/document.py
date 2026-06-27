"""Document-ingestion importer: PDF/DOCX vendor CRM or baseline → candidate BaselineEntry.

Pipeline:
  1. Extract text from each source document (PDF or DOCX).
  2. Build a prompt and send to the configured AI provider.
  3. Parse the JSON response into domain types.
  4. Apply evidence-minimization rules as a hard post-AI enforcement step
     (not just a prompt instruction — the code enforces the invariants so a
     confused model can't accidentally credit a disclaimed control).
  5. Compute and attach the summary roll-up.
  6. Return the candidate BaselineEntry for engineer review.

Nothing is written anywhere — the caller decides what to do with the candidate.

When two documents are supplied (typical: an MSP baseline + the vendor CRM),
the prompt instructs the model to cross-reference them: the baseline resolves
which product is authoritative per control family; the CRM resolves the
per-objective responsibility split. This is the pattern that catches the IA
family correctly — the CRM says "customer IdP owns all of IA", and the baseline
doesn't claim it for this product, so every IA control is customer_owns.
"""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from ..ai.base import AIProvider
from ..baseline import (
    BaselineEntry,
    CandidateState,
    Classification,
    ControlEntry,
    EvidenceSpec,
    ProductMeta,
)

# ---------------------------------------------------------------------------
# Classification prompt
# ---------------------------------------------------------------------------

_SYSTEM_PROMPT = """\
You are a CMMC compliance analyst. Your task: read the vendor document(s) \
provided and extract baseline library entries in JSON.

For each CMMC / NIST 800-171 control mentioned, determine the responsibility \
split and assign one classification:

  provider_satisfies  – the vendor/product materially satisfies this control \
for in-scope assets when correctly configured.
  shared              – the product enables this control but the customer must \
configure or operate part of it.
  customer_owns       – the document explicitly says the CUSTOMER (not the \
product) is responsible. DO NOT credit the vendor. Route this control to the \
product that actually owns it.

Evidence-minimization rules — apply these in your output:
  • Only provider_satisfies and shared controls get evidence entries. \
customer_owns controls must have evidence set to [].
  • Where a single artifact satisfies multiple objectives, list it once.
  • Prefer one authoritative export or config dump over multiple screenshots.
  • Evidence type is one of: screenshot | export | document | link.

Batch an entire control family into one entry (with "control" as a list) when \
the document disclaims the whole family in a single statement \
(e.g. "the customer's IdP is responsible for all IA controls").

Return ONLY a JSON object — no markdown fences, no prose, just JSON:

{
  "product": {
    "name": "...",
    "provider": "...",
    "role": "one-sentence description of what this product does in the stack",
    "assumed_config": ["config assumption 1", "..."],
    "source_docs": ["document filename or title"]
  },
  "controls": [
    {
      "control": "AC.L2-3.1.1",
      "objectives": ["a", "b"],
      "classification": "shared",
      "provider_contribution": "What the product provides for this control.",
      "customer_action": "What the customer must do.",
      "evidence": [
        {"artifact": "Portal user/role list", "type": "export", "kb": "KB article"}
      ],
      "candidate_state": "pending_evidence",
      "note": "optional note",
      "scope_note": "optional scope clarification"
    },
    {
      "control": ["IA.L2-3.5.1", "IA.L2-3.5.2"],
      "objectives": [],
      "classification": "customer_owns",
      "evidence": [],
      "candidate_state": "not_satisfied_by_product",
      "note": "Reason the product disclaims these controls."
    }
  ]
}

Do not invent controls not mentioned in the documents. Do not credit the vendor \
for controls it explicitly disclaims or defers to the customer.\
"""


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------


def extract_text(path: str | Path) -> str:
    """Extract plain text from a PDF or DOCX file."""
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(p)
    if suffix in (".docx", ".doc"):
        return _extract_docx(p)
    raise ValueError(
        f"Unsupported document type: {suffix!r}. Supported: .pdf, .docx"
    )


def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError(
            "pypdf not installed. Run: pip install 'wingrc-backend[ai]'"
        ) from exc
    reader = PdfReader(path)
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n".join(pages)


def _extract_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "python-docx not installed. Run: pip install 'wingrc-backend[ai]'"
        ) from exc
    doc = Document(path)
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # CRMs are often tables — extract those too.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# AI response parsing
# ---------------------------------------------------------------------------


def _strip_fences(text: str) -> str:
    """Remove markdown code fences a model may add despite instructions."""
    cleaned = re.sub(r"^```(?:json)?\s*", "", text.strip(), flags=re.MULTILINE)
    return re.sub(r"```\s*$", "", cleaned, flags=re.MULTILINE).strip()


def _parse_ai_json(raw: str) -> dict[str, Any]:
    return json.loads(_strip_fences(raw))


# ---------------------------------------------------------------------------
# Evidence-minimization enforcement
# ---------------------------------------------------------------------------


def _apply_evidence_minimization(entry: ControlEntry) -> ControlEntry:
    """Hard-enforce minimization rules regardless of what the AI produced.

    This is the code-level guarantee that customer_owns controls never carry
    evidence specs or the wrong candidate_state — even if the model ignores
    the prompt instruction.
    """
    if entry.classification == Classification.CUSTOMER_OWNS:
        return ControlEntry(
            control=entry.control,
            classification=entry.classification,
            candidate_state=CandidateState.NOT_SATISFIED_BY_PRODUCT,
            objectives=entry.objectives,
            provider_contribution=None,
            customer_action=None,
            evidence=[],
            note=entry.note,
            scope_note=entry.scope_note,
        )
    return ControlEntry(
        control=entry.control,
        classification=entry.classification,
        candidate_state=CandidateState.PENDING_EVIDENCE,
        objectives=entry.objectives,
        provider_contribution=entry.provider_contribution,
        customer_action=entry.customer_action,
        evidence=entry.evidence,
        note=entry.note,
        scope_note=entry.scope_note,
    )


def _parse_control(raw: dict[str, Any]) -> ControlEntry:
    cls = Classification(raw["classification"])
    evidence = [
        EvidenceSpec(
            artifact=ev["artifact"],
            type=ev["type"],
            kb=ev.get("kb"),
        )
        for ev in raw.get("evidence", [])
    ]
    entry = ControlEntry(
        control=raw["control"],
        classification=cls,
        candidate_state=CandidateState(
            raw.get("candidate_state", "pending_evidence")
        ),
        objectives=[str(o) for o in raw.get("objectives", [])],
        provider_contribution=raw.get("provider_contribution"),
        customer_action=raw.get("customer_action"),
        evidence=evidence,
        note=raw.get("note"),
        scope_note=raw.get("scope_note"),
    )
    return _apply_evidence_minimization(entry)


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------


def ingest_document(
    *paths: str | Path,
    product_key: str,
    ai_provider: AIProvider,
    category: str = "ESP",
    asset_type: str = "SPA",
    framework: str = "NIST 800-171 Rev 2 / CMMC L2",
) -> BaselineEntry:
    """Extract candidate baseline entries from one or more vendor documents.

    Pass two documents when both a CRM and an MSP baseline are available.
    The AI cross-references them: the baseline resolves which product is
    authoritative per control family; the CRM resolves the per-objective
    responsibility split.

    Args:
        *paths:       One or more PDF or DOCX files (CRM, baseline doc, …).
        product_key:  Baseline library slug, e.g. "rocketcyber". Not inferred
                      from the doc — the author assigns this identifier.
        ai_provider:  The configured AIProvider instance.
        category:     CMMC asset category (default "ESP").
        asset_type:   CMMC asset type (default "SPA").
        framework:    Framework version string.

    Returns:
        A candidate BaselineEntry with summary computed. Nothing is persisted.
    """
    doc_texts: list[str] = []
    source_docs: list[str] = []
    for path in paths:
        p = Path(path)
        source_docs.append(p.name)
        doc_texts.append(f"=== Source: {p.name} ===\n{extract_text(p)}")

    combined = "\n\n".join(doc_texts)
    if len(paths) > 1:
        user_msg = (
            "Two documents are provided. Cross-reference them: the baseline "
            "resolves which product is authoritative per control family; the "
            "CRM resolves the per-objective responsibility split.\n\n"
            + combined
        )
    else:
        user_msg = combined

    raw_json = ai_provider.complete(_SYSTEM_PROMPT, user_msg)
    data = _parse_ai_json(raw_json)

    p_data = data["product"]
    product = ProductMeta(
        key=product_key,
        name=p_data.get("name", product_key),
        provider=p_data.get("provider", ""),
        category=category,
        asset_type=asset_type,
        framework=framework,
        role=p_data.get("role", "").strip(),
        assumed_config=list(p_data.get("assumed_config", [])),
        source_docs=source_docs,
    )

    controls = [_parse_control(c) for c in data.get("controls", [])]
    entry = BaselineEntry(product=product, controls=controls)
    entry.summary = entry.compute_summary()
    return entry
